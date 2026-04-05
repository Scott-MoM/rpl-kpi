from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
MANUALS_DIR = ROOT / "docs" / "manuals"
OUTPUT_DIR = MANUALS_DIR / "docx"


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    if "Manual Bullet" not in styles:
        style = styles.add_style("Manual Bullet", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["List Bullet"]
        style.font.name = "Aptos"
        style.font.size = Pt(11)

    if "Manual Number" not in styles:
        style = styles.add_style("Manual Number", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["List Number"]
        style.font.name = "Aptos"
        style.font.size = Pt(11)


def render_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""


def add_image(doc: Document, manual_path: Path, rel_path: str) -> None:
    image_path = (manual_path.parent / rel_path).resolve()
    if image_path.suffix.lower() == ".svg":
        png_candidate = image_path.with_suffix(".png")
        if png_candidate.exists():
            image_path = png_candidate
    if not image_path.exists():
        doc.add_paragraph(f"[Missing image: {rel_path}]")
        return
    doc.add_picture(str(image_path), width=Inches(6.8))


def add_paragraph_with_formatting(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def convert_manual(manual_path: Path) -> Path:
    doc = Document()
    ensure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.start_type = WD_SECTION.NEW_PAGE

    lines = manual_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            render_table(doc, table_lines)
            continue

        if stripped.startswith("!["):
            match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
            if match:
                add_image(doc, manual_path, match.group(1))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            doc.add_heading(title, level=min(level - 1, 4))
            i += 1
            continue

        if stripped.startswith(">"):
            add_paragraph_with_formatting(doc, stripped[1:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_paragraph_with_formatting(doc, re.sub(r"^\d+\.\s+", "", stripped), style="Manual Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_paragraph_with_formatting(doc, stripped[2:], style="Manual Bullet")
            i += 1
            continue

        add_paragraph_with_formatting(doc, stripped)
        i += 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{manual_path.stem}.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    manual_paths = sorted(MANUALS_DIR.glob("*.md"))
    for manual_path in manual_paths:
        output_path = convert_manual(manual_path)
        print(output_path)


if __name__ == "__main__":
    main()
